"""
Report Generator Module

Generates reports from code analysis issues in PDF, DOC, and MD formats.
"""

import re
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any

logger = logging.getLogger(__name__)


def detect_formats(prompt: str) -> List[str]:
    """
    Detect format keywords in the prompt.
    
    Looks for patterns like:
    - "generate pdf", "create pdf", "export pdf"
    - "in pdf format", "as pdf", "pdf report"
    - Standalone mentions: "pdf", "doc", "md"
    
    Args:
        prompt: The user prompt
        
    Returns:
        List of detected formats (lowercase): ['pdf', 'doc', 'md']
    """
    prompt_lower = prompt.lower()
    detected = []
    
    # Pattern 1: Action words followed by format
    # e.g., "generate pdf", "create doc", "export md"
    action_pattern = r'(?:generate|create|export|save|make|build|produce|output|write)\s+(?:a\s+)?(pdf|doc|docx|md|markdown)'
    matches = re.findall(action_pattern, prompt_lower)
    for match in matches:
        if match == 'docx':
            detected.append('doc')
        elif match == 'markdown':
            detected.append('md')
        else:
            detected.append(match)
    
    # Pattern 2: Format in context
    # e.g., "in pdf format", "as pdf", "pdf report", "pdf file"
    context_pattern = r'(?:in|as|for|to)\s+(?:a\s+)?(pdf|doc|docx|md|markdown)\s+(?:format|report|file|document)?'
    matches = re.findall(context_pattern, prompt_lower)
    for match in matches:
        if match == 'docx':
            detected.append('doc')
        elif match == 'markdown':
            detected.append('md')
        else:
            detected.append(match)
    
    # Pattern 3: Standalone mentions
    # e.g., "pdf", "doc", "md"
    standalone_pattern = r'\b(pdf|doc|docx|md|markdown)\b'
    matches = re.findall(standalone_pattern, prompt_lower)
    for match in matches:
        if match == 'docx':
            detected.append('doc')
        elif match == 'markdown':
            detected.append('md')
        else:
            detected.append(match)
    
    # Remove duplicates and return
    return list(set(detected))


def generate_report_summary(
    prompt: str,
    context: str,
    model: Optional[str] = None
) -> str:
    """
    Generate a report summary using LLM.
    
    Args:
        prompt: User's report generation prompt
        context: Context string with issues and summary
        model: Optional model name to use
        
    Returns:
        Generated report summary text
    """
    try:
        from langchain_core.prompts import ChatPromptTemplate
        from config import get_llm, get_settings
        
        settings = get_settings()
        
        if not settings.use_llm_analysis:
            raise ValueError("LLM analysis is not enabled")
        
        llm = get_llm(model_override=model)
        
        system_prompt = """You are a technical report generator. You create comprehensive, well-structured reports 
about code analysis issues. Your reports should be professional, clear, and actionable.

Use the provided context about code issues to generate a detailed report. Structure your report with:
- Clear headings and sections
- Bullet points for lists
- Code blocks where appropriate (using markdown syntax)
- Prioritized recommendations

Format the output as markdown."""
        
        chat_prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "Context:\n{context}\n\nUser Request:\n{prompt}\n\nGenerate a comprehensive report:")
        ])
        
        messages = chat_prompt.format_messages(
            context=context,
            prompt=prompt
        )
        
        response = llm.invoke(messages)
        return response.content
        
    except Exception as e:
        logger.error(f"Failed to generate report summary: {e}")
        raise


def generate_markdown(summary: str, issues: List[Dict[str, Any]], summary_stats: Dict[str, Any]) -> str:
    """
    Format summary and issues into a complete markdown report.
    
    Args:
        summary: LLM-generated summary
        issues: List of issue dictionaries
        summary_stats: Summary statistics dict
        
    Returns:
        Complete markdown report content
    """
    md = f"""# Code Analysis Report

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Executive Summary

{summary}

## Statistics

- **Total Issues:** {summary_stats.get('total', 0)}
- **By Type:**
  - Security: {summary_stats.get('by_type', {}).get('security', 0)}
  - Performance: {summary_stats.get('by_type', {}).get('performance', 0)}
  - Architecture: {summary_stats.get('by_type', {}).get('architecture', 0)}
- **By Risk Level:**
  - Critical: {summary_stats.get('by_risk_level', {}).get('critical', 0)}
  - High: {summary_stats.get('by_risk_level', {}).get('high', 0)}
  - Medium: {summary_stats.get('by_risk_level', {}).get('medium', 0)}
  - Low: {summary_stats.get('by_risk_level', {}).get('low', 0)}

## Issues by Risk Level

"""
    
    # Group issues by risk level
    by_risk = {
        'critical': [],
        'high': [],
        'medium': [],
        'low': []
    }
    
    for issue in issues:
        risk = issue.get('risk_level', 'low')
        if risk in by_risk:
            by_risk[risk].append(issue)
    
    for risk_level in ['critical', 'high', 'medium', 'low']:
        risk_issues = by_risk[risk_level]
        if risk_issues:
            md += f"\n### {risk_level.capitalize()} Risk Issues\n\n"
            for issue in risk_issues[:20]:  # Limit to top 20 per category
                md += f"#### {issue.get('title', 'Untitled')}\n\n"
                md += f"**Location:** `{issue.get('location', 'Unknown')}`\n\n"
                md += f"**Type:** {issue.get('type', 'unknown').capitalize()}\n\n"
                md += f"**Description:**\n{issue.get('description', 'No description')}\n\n"
                
                if issue.get('code_snippet'):
                    md += f"**Code Snippet:**\n```python\n{issue.get('code_snippet')}\n```\n\n"
                
                if issue.get('solution'):
                    md += f"**Solution:**\n{issue.get('solution')}\n\n"
                
                md += "---\n\n"
    
    md += "\n## Recommendations\n\n"
    md += "1. Address critical and high-risk issues immediately\n"
    md += "2. Review and refactor code with multiple issues\n"
    md += "3. Implement automated testing to prevent regressions\n"
    md += "4. Establish code review processes\n"
    md += "5. Monitor and track issue resolution progress\n"
    
    return md


def save_markdown(content: str, output_path: Path) -> Path:
    """
    Save markdown content to a file.
    
    Args:
        content: Markdown content
        output_path: Path to save the file
        
    Returns:
        Path to saved file
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding='utf-8')
    return output_path


def generate_pdf(markdown_content: str, output_path: Path) -> Path:
    """
    Convert markdown to PDF using weasyprint.
    
    Args:
        markdown_content: Markdown content to convert
        output_path: Path for output PDF file
        
    Returns:
        Path to generated PDF file
    """
    try:
        import markdown
        from weasyprint import HTML
        
        # Try to use extensions, fall back to basic if not available
        extensions = ['fenced_code']
        
        # Convert markdown to HTML
        # If extensions fail, try without them
        try:
            html_content = markdown.markdown(
                markdown_content,
                extensions=['codehilite', 'fenced_code', 'tables']
            )
        except Exception:
            # Fall back to basic markdown if extensions aren't available
            html_content = markdown.markdown(
                markdown_content,
                extensions=['fenced_code']
            )
        
        # Wrap in HTML document structure
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: 'Helvetica', 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #2563eb; border-bottom: 2px solid #2563eb; padding-bottom: 0.3em; }}
        h2 {{ color: #1e40af; margin-top: 1.5em; }}
        h3 {{ color: #1e3a8a; margin-top: 1.2em; }}
        h4 {{ color: #1e3a8a; margin-top: 1em; }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 1em;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #2563eb;
            color: white;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        # Generate PDF
        output_path.parent.mkdir(parents=True, exist_ok=True)
        HTML(string=full_html).write_pdf(output_path)
        
        return output_path
        
    except ImportError as e:
        error_msg = (
            "weasyprint is not installed. To install:\n"
            "1. Install system dependencies (macOS): brew install cairo pango gdk-pixbuf libffi\n"
            "2. Install weasyprint: pip install weasyprint\n"
            f"Original error: {str(e)}"
        )
        logger.error(error_msg)
        raise ImportError(error_msg) from e
    except Exception as e:
        logger.error(f"Failed to generate PDF: {e}")
        raise


def generate_doc(markdown_content: str, output_path: Path) -> Path:
    """
    Convert markdown to DOCX using python-docx.
    
    Args:
        markdown_content: Markdown content to convert
        output_path: Path for output DOCX file
        
    Returns:
        Path to generated DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Parse markdown and convert to docx
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            # Code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1  # Skip the opening ```
                while i < len(lines):
                    if lines[i].strip().startswith('```'):
                        i += 1  # Skip the closing ```
                        break
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.style = 'Intense Quote'
                # Skip the increment at the end since we've already handled i
                continue
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered lists
            elif re.match(r'^\d+\.\s+', line):
                doc.add_paragraph(re.sub(r'^\d+\.\s+', '', line), style='List Number')
            # Bold text
            elif '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            # Regular paragraph
            else:
                # Handle inline formatting
                p = doc.add_paragraph()
                # Simple text for now - can be enhanced later
                p.add_run(line)
            
            i += 1
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return output_path
        
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Failed to generate DOCX: {e}")
        raise
